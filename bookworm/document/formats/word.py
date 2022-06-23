# coding: utf-8

from __future__ import annotations

import contextlib
import os
import subprocess
from io import BytesIO
from pathlib import Path

import mammoth
import msoffcrypto
import msoffcrypto.exceptions
from diskcache import Cache
from docx import Document as DocxDocumentReader
from selectolax.parser import HTMLParser

from bookworm import app
from bookworm.concurrency import process_worker, threaded_worker
from bookworm.document.uri import DocumentUri
from bookworm.logger import logger
from bookworm.paths import app_path, home_data_path
from bookworm.utils import NEWLINE, escape_html, generate_file_md5

from .. import ChangeDocument
from .. import DocumentCapability as DC
from .. import DocumentEncryptedError, DocumentError, DummyDocument
from .html import BaseHtmlDocument

log = logger.getChild(__name__)
TAGS_TO_UNWRAP = []
TAGS_TO_REMOVE = [
    "img",
    "style",
]


class WordDocument(BaseHtmlDocument):

    format = "docx"
    # Translators: the name of a document file format
    name = _("Word Document")
    extensions = ("*.docx",)

    def read(self):
        file_bytes = self.get_file_system_path().read_bytes()
        data_buf = BytesIO(file_bytes)
        is_encrypted_document = self.is_encrypted(data_buf)
        if is_encrypted_document:
            if (decryption_key := self.uri.view_args.get("decryption_key")) is not None:
                self.try_decrypt(data_buf, decryption_key)
            else:
                raise DocumentEncryptedError(self)
        self.__html_content = self._get_html_content_from_docx(data_buf, is_encrypted_document)
        super().read()

    def get_html(self):
        return self.__html_content

    def parse_html(self):
        return self.parse_to_full_text()

    def is_encrypted(self, data_buf):
        data_buf.seek(0)
        first_2000b = data_buf.read(2000).replace(b"\0", b" ")
        return b"E n c r y p t e d P a c k a g e" in first_2000b

    def try_decrypt(self, data_buf, decryption_key):
        data_buf.seek(0)
        msc_file = msoffcrypto.OfficeFile(data_buf)
        msc_file.load_key(password=decryption_key)
        try:
            out_buf = BytesIO()
            msc_file.decrypt(out_buf, True)
        except msoffcrypto.exceptions.InvalidKeyError:
            raise DocumentEncryptedError(self)
        else:
            with contextlib.closing(out_buf):
                data_buf.seek(0)
                data_buf.write(out_buf.getvalue())

    def _get_html_content_from_docx(self, data_buf, is_encrypted_document):
        data_buf.seek(0)
        cache = Cache(self._get_cache_directory(), eviction_policy="least-frequently-used")
        cache_key = self.uri.to_uri_string()
        if cached_html_content := cache.get(cache_key):
            return cached_html_content.decode("utf-8")
        result = mammoth.convert_to_html(data_buf, include_embedded_style_map=False)
        data_buf.seek(0)
        html_content = self.make_proper_html(result.value, data_buf)
        if not is_encrypted_document:
            cache.set(cache_key, html_content.encode("utf-8"))
        return html_content

    def make_proper_html(self, html_string, data_buf):
        parsed = HTMLParser(html_string)
        parsed.body.unwrap_tags(TAGS_TO_UNWRAP)
        parsed.strip_tags(TAGS_TO_REMOVE)
        html_string = parsed.body.html
        docx = DocxDocumentReader(data_buf)
        props = docx.core_properties
        doc_title = props.title.strip()
        if not doc_title or doc_title.lower() == "word document":
            doc_title = self.get_file_system_path().stem.strip()
        doc_author = escape_html(props.author or "")
        data_buf.close()
        return NEWLINE.join(
            [
                "<!DOCTYPE html>",
                "<html>",
                "<head>",
                '<meta charset="utf-8"/>',
                f'<meta name="author" content="{doc_author}"/>',
                f"<title>{escape_html(doc_title)}</title>",
                "</head>",
                html_string,
                "</html>",
            ]
        )

    def _get_cache_directory(self):
        return os.fspath(home_data_path(".docx_to_html"))


class Word97Document(DummyDocument):

    format = "doc"
    # Translators: the name of a document file format
    name = _("Word 97 - 2003 Document")
    extensions = ("*.doc",)
    capabilities = DC.ASYNC_READ

    def read(self):
        converted_file = threaded_worker.submit(
            self.get_converted_filename, self.get_file_system_path()
        ).result()
        raise ChangeDocument(
            old_uri=self.uri,
            new_uri=DocumentUri.from_filename(converted_file),
            reason="Doc converted to docbook",
        )

    @classmethod
    def get_converted_filename(cls, filename):
        storage_area = cls.get_storage_area()
        target_file = storage_area / f"{generate_file_md5(filename)}.docbook"
        if not target_file.exists():
            docbook_content = cls.convert_to_docbook(filename)
            target_file.write_bytes(docbook_content)
        return target_file

    @classmethod
    def convert_to_docbook(cls, filename):
        args = [cls.get_antiword_executable_path(), "-x", "db", filename]
        creationflags = subprocess.CREATE_NO_WINDOW
        startupinfo = subprocess.STARTUPINFO()
        startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        ret = subprocess.run(
            args,
            capture_output=True,
            creationflags=creationflags,
            startupinfo=startupinfo,
        )
        return ret.stdout

    @staticmethod
    def get_storage_area():
        storage_area = home_data_path(".doc_to_docbook")
        storage_area.mkdir(parents=True, exist_ok=True)
        return storage_area

    @staticmethod
    def get_antiword_executable_path():
        if app.is_frozen:
            return app_path("antiword", "bin", "antiword.exe")
        else:
            return (
                Path.cwd()
                / "scripts"
                / "executables"
                / "antiword"
                / "bin"
                / "antiword.exe"
            )
